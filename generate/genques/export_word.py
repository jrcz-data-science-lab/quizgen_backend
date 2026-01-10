from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO


def export_quiz_to_word(quiz_data):
    """
    Export quiz data to a Word document.
    
    Args:
        quiz_data: Dictionary containing quiz with 'mcq' array
        
    Returns:
        BytesIO object containing the Word document
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Quiz - Multiple Choice Questions', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Add spacing
    
    # Get the MCQ array
    mcqs = quiz_data.get('mcq', [])
    
    if not mcqs:
        doc.add_paragraph('No questions available.')
    
    # Add each question
    for idx, question in enumerate(mcqs, 1):
        # Question number and text
        q_paragraph = doc.add_paragraph()
        q_run = q_paragraph.add_run(f"Question {idx}: {question.get('question', '')}")
        q_run.bold = True
        q_run.font.size = Pt(12)
        
        # Add options
        options = question.get('options', [])
        for option in options:
            option_paragraph = doc.add_paragraph(option, style='List Bullet')
            option_paragraph.paragraph_format.left_indent = Pt(20)
        
        # Add correct answer
        answer_paragraph = doc.add_paragraph()
        answer_run = answer_paragraph.add_run(f"Correct Answer: {question.get('answer', '')}")
        answer_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
        answer_run.bold = True
        
        # Add spacing between questions
        doc.add_paragraph()
    
    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
